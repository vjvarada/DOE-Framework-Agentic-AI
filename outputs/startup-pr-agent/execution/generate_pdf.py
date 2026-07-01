#!/usr/bin/env python3
"""
PDF and Document Generator for HR documents.

Generates professional PDFs (offer letters, JDs, evaluation reports, scorecards)
and can also create Word documents (.docx).

Usage:
    python generate_pdf.py --type job-description --data .tmp/jds/role.json --output .tmp/documents/
    python generate_pdf.py --type offer-letter --data .tmp/offer_data.json --output .tmp/documents/
    python generate_pdf.py --type evaluation-report --data .tmp/evaluations/eval_candidate.json --output .tmp/documents/
    python generate_pdf.py --type interview-scorecard --data .tmp/scorecard_data.json --output .tmp/documents/
    python generate_pdf.py --type custom --title "Document Title" --content-file .tmp/content.md --output .tmp/documents/
    python generate_pdf.py --type job-description --data .tmp/jds/role.json --format docx --output .tmp/documents/
"""

import os
import sys
import json
import argparse
from pathlib import Path
from datetime import datetime


def _sanitize(text: str) -> str:
    """Replace Unicode characters that aren't supported by standard PDF fonts."""
    replacements = {
        '\u2014': '-',   # em dash
        '\u2013': '-',   # en dash
        '\u2018': "'",   # left single quote
        '\u2019': "'",   # right single quote
        '\u201c': '"',   # left double quote
        '\u201d': '"',   # right double quote
        '\u2022': '-',   # bullet
        '\u2713': '[Y]', # checkmark
        '\u2717': '[X]', # cross mark
        '\u2026': '...', # ellipsis
        '\u00b7': '-',   # middle dot
        '\u25aa': '-',   # small black square
        '\u25ba': '>',   # triangle
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Fallback: replace any remaining non-latin-1 chars
    return text.encode('latin-1', errors='replace').decode('latin-1')


def generate_jd_pdf(data: dict, output_path: Path):
    """Generate a job description PDF."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Header
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, data.get("title", "Job Description"), ln=True, align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, _sanitize(f"Fracktal Works Pvt. Ltd - {data.get('department', '')}"), ln=True, align="C")
    pdf.cell(0, 8, f"{data.get('location', '')} | {data.get('work_type', '')}", ln=True, align="C")
    pdf.ln(8)

    # Company description
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "About Fracktal Works", ln=True)
    pdf.set_font("Helvetica", "", 10)
    desc = data.get("company_description", "")
    pdf.multi_cell(0, 5, _sanitize(desc))
    pdf.ln(5)

    # Role summary
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Role Summary", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 5,
                   f"Level: {data.get('level', 'N/A')} | "
                   f"Experience: {data.get('experience_required', 'N/A')} | "
                   f"Education: {data.get('education_required', 'N/A')}")
    pdf.ln(3)

    # Responsibilities
    if data.get("responsibilities"):
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Key Responsibilities", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for r in data["responsibilities"]:
            pdf.multi_cell(0, 5, _sanitize(f"  *  {r}"))
        pdf.ln(3)

    # Skills
    if data.get("required_skills"):
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Required Skills", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for s in data["required_skills"]:
            pdf.cell(0, 5, _sanitize(f"  *  {s}"), ln=True)
        pdf.ln(3)

    # Benefits
    if data.get("benefits"):
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "What We Offer", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for b in data["benefits"]:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(0, 5, _sanitize(f"  *  {b}"))
        pdf.ln(3)

    # Footer
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 8, f"Fracktal Works is an equal opportunity employer. Generated {datetime.now().strftime('%Y-%m-%d')}",
             ln=True, align="C")

    pdf.output(str(output_path))
    print(f"  Generated JD PDF: {output_path}")


def generate_offer_letter_pdf(data: dict, output_path: Path):
    """Generate an offer letter PDF."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Header
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "FRACKTAL WORKS PVT. LTD", ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "Bangalore, India | www.fracktal.in", ln=True, align="C")
    pdf.ln(5)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)

    # Date
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 6, f"Date: {data.get('date', datetime.now().strftime('%B %d, %Y'))}", ln=True)
    pdf.ln(3)

    # To
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 6, f"To: {data.get('candidate_name', '[Candidate Name]')}", ln=True)
    if data.get("candidate_address"):
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, data["candidate_address"])
    pdf.ln(5)

    # Subject
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, _sanitize(f"Subject: Offer of Employment - {data.get('title', '[Position]')}"), ln=True)
    pdf.ln(5)

    # Body
    pdf.set_font("Helvetica", "", 11)
    body = f"Dear {data.get('candidate_name', '[Candidate Name]')},\n\n"
    body += f"We are pleased to offer you the position of {data.get('title', '[Position]')} "
    body += f"in the {data.get('department', '[Department]')} department at Fracktal Works Pvt. Ltd.\n\n"

    body += "Terms of Employment:\n\n"
    body += f"  Position: {data.get('title', '[Position]')}\n"
    body += f"  Department: {data.get('department', '[Department]')}\n"
    body += f"  Reporting To: {data.get('reporting_to', '[Manager Name]')}\n"
    body += f"  Start Date: {data.get('start_date', '[Start Date]')}\n"
    body += f"  Location: {data.get('location', 'Bangalore, India')}\n"
    body += f"  Employment Type: {data.get('employment_type', 'Full-Time')}\n"

    if data.get("ctc"):
        body += f"  Annual CTC: {data['ctc']}\n"
    if data.get("probation_period"):
        body += f"  Probation Period: {data['probation_period']}\n"
    if data.get("notice_period"):
        body += f"  Notice Period: {data['notice_period']}\n"

    body += "\nThis offer is contingent upon satisfactory completion of background verification "
    body += "and submission of all required documents.\n\n"
    body += "Please confirm your acceptance by signing and returning this letter "
    body += f"by {data.get('acceptance_deadline', '[Deadline]')}.\n\n"
    body += "We look forward to welcoming you to the Fracktal Works team.\n\n"
    body += "Warm regards,\n\n\n"
    body += f"{data.get('signatory_name', '[Authorized Signatory]')}\n"
    body += f"{data.get('signatory_title', '[Title]')}\n"
    body += "Fracktal Works Pvt. Ltd\n\n"

    body += "---\n\n"
    body += "ACCEPTANCE\n\n"
    body += f"I, {data.get('candidate_name', '[Candidate Name]')}, accept the offer of employment "
    body += "as described above.\n\n"
    body += "Signature: ____________________    Date: ____________________"

    pdf.multi_cell(0, 6, _sanitize(body))

    pdf.output(str(output_path))
    print(f"  Generated Offer Letter PDF: {output_path}")


def generate_evaluation_pdf(data: dict, output_path: Path):
    """Generate a candidate evaluation report PDF."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Header
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "CANDIDATE EVALUATION REPORT", ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "Fracktal Works Pvt. Ltd - Confidential", ln=True, align="C")
    pdf.ln(5)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    c = data.get("candidate", {})
    r = data.get("role", {})
    s = data.get("scores", {})

    # Candidate info
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Candidate Information", ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 6, f"Name: {c.get('name', 'N/A')}", ln=True)
    pdf.cell(0, 6, f"Email: {c.get('email', 'N/A')}", ln=True)
    pdf.cell(0, 6, f"Resume: {c.get('file', 'N/A')}", ln=True)
    pdf.ln(3)

    # Position
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Position Details", ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 6, f"Title: {r.get('title', 'N/A')}", ln=True)
    pdf.cell(0, 6, f"Department: {r.get('department', 'N/A')}", ln=True)
    pdf.cell(0, 6, f"Level: {r.get('level', 'N/A')}", ln=True)
    pdf.ln(3)

    # Scores
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 8, "Evaluation Scores", ln=True)
    pdf.set_font("Helvetica", "", 11)

    skills = s.get("skills", {})
    pdf.cell(0, 6, f"Skills Match: {skills.get('score', 'N/A')}/10 ({skills.get('match_pct', 'N/A')}% overlap)", ln=True)
    exp = s.get("experience", {})
    pdf.cell(0, 6, _sanitize(f"Experience: {exp.get('score', 'N/A')}/10 - {exp.get('note', '')}"), ln=True)
    edu = s.get("education", {})
    pdf.cell(0, 6, _sanitize(f"Education: {edu.get('score', 'N/A')}/10 - {edu.get('note', '')}"), ln=True)
    ind = s.get("industry_fit", {})
    pdf.cell(0, 6, f"Industry Fit: {ind.get('score', 'N/A')}/5", ln=True)
    pdf.ln(3)

    # Overall
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, f"Overall Score: {s.get('overall', 'N/A')}/10", ln=True)
    pdf.set_font("Helvetica", "B", 14)
    recommendation = data.get("recommendation", "N/A")
    pdf.cell(0, 10, f"Recommendation: {recommendation}", ln=True)
    pdf.ln(3)

    # Details
    if skills.get("matched"):
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Skills Matched:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for m in skills["matched"]:
            pdf.cell(0, 5, _sanitize(f"  [Y] {m['required']}"), ln=True)

    if skills.get("missing"):
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Skills Missing:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        for m in skills["missing"]:
            pdf.cell(0, 5, _sanitize(f"  [X] {m}"), ln=True)

    if ind.get("signals"):
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, "Industry Signals:", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 5, _sanitize(", ".join(ind["signals"][:15])))

    # Footer
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 8)
    pdf.cell(0, 5, f"Evaluated: {data.get('evaluated_at', datetime.now().isoformat())[:10]}", ln=True, align="C")
    pdf.cell(0, 5, "This document is confidential. For internal use only.", ln=True, align="C")

    pdf.output(str(output_path))
    print(f"  Generated Evaluation PDF: {output_path}")


def generate_scorecard_pdf(data: dict, output_path: Path):
    """Generate an interview scorecard PDF."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "INTERVIEW SCORECARD", ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "Fracktal Works Pvt. Ltd - Confidential", ln=True, align="C")
    pdf.ln(5)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 6, f"Candidate: {data.get('candidate_name', '____________________')}", ln=True)
    pdf.cell(0, 6, f"Position: {data.get('position', '____________________')}", ln=True)
    pdf.cell(0, 6, f"Interviewer: {data.get('interviewer', '____________________')}", ln=True)
    pdf.cell(0, 6, f"Date: {data.get('date', '____________________')}", ln=True)
    pdf.cell(0, 6, f"Round: {data.get('round', '____________________')}", ln=True)
    pdf.ln(5)

    # Criteria
    criteria = data.get("criteria", [
        "Technical Skills", "Problem Solving", "Communication",
        "Culture Fit", "Domain Knowledge", "Leadership/Initiative"
    ])

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(80, 8, "Criteria", border=1)
    pdf.cell(20, 8, "Score", border=1, align="C")
    pdf.cell(90, 8, "Notes", border=1)
    pdf.ln()

    pdf.set_font("Helvetica", "", 10)
    for c in criteria:
        pdf.cell(80, 12, f"  {c}", border=1)
        pdf.cell(20, 12, "    /5", border=1, align="C")
        pdf.cell(90, 12, "", border=1)
        pdf.ln()

    pdf.ln(5)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Overall Assessment:", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "[ ] Strong Hire  [ ] Hire  [ ] Maybe  [ ] No Hire", ln=True)
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Key Strengths:", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 12, "")
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Concerns:", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 12, "")
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Additional Notes:", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 20, "")

    pdf.ln(10)
    pdf.cell(0, 6, "Interviewer Signature: ____________________", ln=True)

    pdf.output(str(output_path))
    print(f"  Generated Scorecard PDF: {output_path}")


def generate_custom_pdf(title: str, content: str, output_path: Path):
    """Generate a custom PDF from text/markdown content."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, title, ln=True, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, "Fracktal Works Pvt. Ltd", ln=True, align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    # Simple markdown-to-pdf: handle headers and bullet points
    for line in content.split('\n'):
        stripped = line.strip()
        if stripped.startswith('# '):
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, stripped[2:], ln=True)
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith('## '):
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 9, stripped[3:], ln=True)
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith('### '):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, stripped[4:], ln=True)
            pdf.set_font("Helvetica", "", 11)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            pdf.multi_cell(0, 5, f"  \u2022  {stripped[2:]}")
        elif stripped.startswith('---'):
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(3)
        elif stripped:
            pdf.multi_cell(0, 5, stripped)
        else:
            pdf.ln(3)

    pdf.output(str(output_path))
    print(f"  Generated custom PDF: {output_path}")


def generate_docx(title: str, content: str, output_path: Path):
    """Generate a Word document from text content."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Fracktal Works Pvt. Ltd")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # spacer

    for line in content.split('\n'):
        stripped = line.strip()
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('---'):
            pass  # Skip horizontal rules
        elif stripped:
            doc.add_paragraph(stripped)

    doc.save(str(output_path))
    print(f"  Generated DOCX: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate HR documents as PDF or DOCX")
    parser.add_argument("--type", required=True,
                        choices=["job-description", "offer-letter", "evaluation-report",
                                 "interview-scorecard", "custom"],
                        help="Document type to generate")
    parser.add_argument("--data", help="Path to JSON data file")
    parser.add_argument("--title", help="Document title (for custom type)")
    parser.add_argument("--content-file", help="Path to content file (for custom type)")
    parser.add_argument("--format", default="pdf", choices=["pdf", "docx"],
                        help="Output format")
    parser.add_argument("--output", default=".tmp/documents/", help="Output directory")
    args = parser.parse_args()

    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    if args.type == "custom":
        if not args.content_file and not args.title:
            print("ERROR: --content-file or --title required for custom type")
            sys.exit(1)
        title = args.title or "Document"
        if args.content_file:
            with open(args.content_file, "r", encoding="utf-8") as f:
                content = f.read()
        else:
            content = ""
        slug = title.lower().replace(" ", "-")[:40]
        if args.format == "docx":
            out_path = output_dir / f"{slug}.docx"
            generate_docx(title, content, out_path)
        else:
            out_path = output_dir / f"{slug}.pdf"
            generate_custom_pdf(title, content, out_path)
        return

    if not args.data:
        print("ERROR: --data required for this document type")
        sys.exit(1)

    with open(args.data, "r", encoding="utf-8") as f:
        data = json.load(f)

    if args.type == "job-description":
        slug = data.get("title", "job-description").lower().replace(" ", "-")
        if args.format == "docx":
            from io import StringIO
            # For DOCX, convert JD data to markdown first, then to docx
            # Quick approach: format as text then generate
            lines = [f"# {data.get('title', 'Job Description')}"]
            lines.append(f"\n**Department:** {data.get('department', '')}")
            lines.append(f"**Level:** {data.get('level', '')}")
            lines.append(f"**Location:** {data.get('location', '')}")
            lines.append(f"\n## About Fracktal Works\n{data.get('company_description', '')}")
            if data.get("responsibilities"):
                lines.append("\n## Key Responsibilities")
                for r in data["responsibilities"]:
                    lines.append(f"- {r}")
            if data.get("required_skills"):
                lines.append("\n## Required Skills")
                for s in data["required_skills"]:
                    lines.append(f"- {s}")
            content = "\n".join(lines)
            out_path = output_dir / f"{slug}.docx"
            generate_docx(data.get("title", "Job Description"), content, out_path)
        else:
            out_path = output_dir / f"{slug}.pdf"
            generate_jd_pdf(data, out_path)

    elif args.type == "offer-letter":
        slug = data.get("candidate_name", "offer").lower().replace(" ", "-")
        out_path = output_dir / f"offer-{slug}.pdf"
        generate_offer_letter_pdf(data, out_path)

    elif args.type == "evaluation-report":
        name = data.get("candidate", {}).get("name", "candidate")
        slug = name.lower().replace(" ", "-")
        out_path = output_dir / f"eval-{slug}.pdf"
        generate_evaluation_pdf(data, out_path)

    elif args.type == "interview-scorecard":
        name = data.get("candidate_name", "scorecard")
        slug = name.lower().replace(" ", "-") if name != "scorecard" else "scorecard"
        out_path = output_dir / f"scorecard-{slug}.pdf"
        generate_scorecard_pdf(data, out_path)


if __name__ == "__main__":
    main()
